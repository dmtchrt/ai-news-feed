from __future__ import annotations

from dataclasses import dataclass
from datetime import UTC, datetime
from pathlib import Path
from shutil import copyfile
from typing import Any

import httpx
import pytest
from docx import Document

from ai_news_feed.domain.models import CollectionCursor
from ai_news_feed.extraction import ContentExtractor, ExtractedItem
from ai_news_feed.sources.presets import expertosphere_source, telegram_preview_source
from ai_news_feed.sources.telegram import TelegramWebPreviewConnector, TelethonConnector


@pytest.mark.asyncio
async def test_web_preview_collects_normal_channel_and_cursor(fixture_dir: Path) -> None:
    html = (fixture_dir / "telegram" / "ailev_preview.html").read_text(encoding="utf-8")

    def handler(request: httpx.Request) -> httpx.Response:
        return httpx.Response(200, text=html, request=request)

    async with httpx.AsyncClient(transport=httpx.MockTransport(handler)) as client:
        connector = TelegramWebPreviewConnector(client)
        first = await connector.collect(telegram_preview_source("@ailev_blog"))
        second = await connector.collect(
            telegram_preview_source("@ailev_blog"),
            first.next_cursor,
        )

    assert [item.external_id for item in first.raw_items] == ["9001", "9002"]
    assert first.raw_items[0].raw_text == "Системное мышление и ИИ\nПолный текст обычного поста."
    assert not second.raw_items


@pytest.mark.asyncio
async def test_empty_expertosphere_preview_does_not_silently_fallback(
    fixture_dir: Path,
) -> None:
    html = (fixture_dir / "telegram" / "expertosphere_empty.html").read_text(encoding="utf-8")

    def handler(request: httpx.Request) -> httpx.Response:
        return httpx.Response(200, text=html, request=request)

    config = telegram_preview_source("@expertosphere")
    async with httpx.AsyncClient(transport=httpx.MockTransport(handler)) as client:
        batch = await TelegramWebPreviewConnector(client).collect(config)

    assert not batch.raw_items
    assert batch.errors[0].code == "preview_unavailable"
    assert expertosphere_source().collector.value == "telethon"


@pytest.mark.asyncio
async def test_preview_document_has_no_fake_download_reference(fixture_dir: Path) -> None:
    html = (fixture_dir / "telegram" / "document_preview.html").read_text(encoding="utf-8")

    def handler(request: httpx.Request) -> httpx.Response:
        return httpx.Response(200, text=html, request=request)

    async with httpx.AsyncClient(transport=httpx.MockTransport(handler)) as client:
        batch = await TelegramWebPreviewConnector(client).collect(
            telegram_preview_source("@reports")
        )

    attachment = batch.raw_items[0].attachments[0]
    assert attachment.name == "AI-market-2026.pdf"
    assert attachment.size == int(2.5 * 1024**2)
    assert attachment.download_ref is None


@dataclass
class _FakeFile:
    name: str
    mime_type: str
    size: int
    ext: str


@dataclass
class _FakeMessage:
    id: int
    date: datetime
    message: str
    document: object | None = None
    file: _FakeFile | None = None
    views: int | None = None
    forwards: int | None = None
    grouped_id: int | None = None


class _AsyncMessages:
    def __init__(self, messages: list[_FakeMessage]) -> None:
        self._messages = messages

    def __aiter__(self) -> _AsyncMessages:
        self._iterator = iter(self._messages)
        return self

    async def __anext__(self) -> _FakeMessage:
        try:
            return next(self._iterator)
        except StopIteration as exc:
            raise StopAsyncIteration from exc


class _FakeTelethonClient:
    def __init__(self, messages: list[_FakeMessage], document_path: Path) -> None:
        self.messages = messages
        self.document_path = document_path
        self.iter_call: dict[str, Any] = {}

    def iter_messages(self, handle: str, **kwargs: Any) -> _AsyncMessages:
        self.iter_call = {"handle": handle, **kwargs}
        min_id = int(kwargs["min_id"])
        return _AsyncMessages([message for message in self.messages if message.id > min_id])

    async def download_media(self, message: _FakeMessage, file: str) -> str:
        copyfile(self.document_path, file)
        return file


@pytest.mark.asyncio
async def test_telethon_downloads_expertosphere_docx_and_extractor_reads_it(
    tmp_path: Path,
) -> None:
    report_path = tmp_path / "market-report.docx"
    report = Document()
    report.add_heading("Рынок искусственного интеллекта", level=1)
    report.add_paragraph("Объём рынка вырос, а компании ускорили внедрение ИИ.")
    report.save(str(report_path))

    message = _FakeMessage(
        id=501,
        date=datetime(2026, 8, 31, 7, 0, tzinfo=UTC),
        message="Новый отраслевой отчёт",
        document=object(),
        file=_FakeFile(
            name="market-report.docx",
            mime_type=("application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
            size=report_path.stat().st_size,
            ext=".docx",
        ),
        views=1200,
    )
    client = _FakeTelethonClient([message], report_path)
    connector = TelethonConnector(client, download_dir=tmp_path / "downloads")
    batch = await connector.collect(expertosphere_source(), CollectionCursor(message_id=500))

    assert batch.next_cursor == CollectionCursor(message_id=501)
    assert client.iter_call["handle"] == "expertosphere"
    attachment = batch.raw_items[0].attachments[0]
    assert attachment.download_ref is not None
    assert Path(attachment.download_ref).is_file()
    extracted = ContentExtractor().extract(batch.raw_items[0])
    assert isinstance(extracted, ExtractedItem)
    assert "Новый отраслевой отчёт" in extracted.text
    assert "Объём рынка вырос" in extracted.text


class _FailingDownloadTelethonClient:
    def __init__(self, messages: list[_FakeMessage]) -> None:
        self.messages = messages

    def iter_messages(self, handle: str, **kwargs: Any) -> _AsyncMessages:
        min_id = int(kwargs["min_id"])
        return _AsyncMessages([message for message in self.messages if message.id > min_id])

    async def download_media(self, message: _FakeMessage, file: str) -> str:
        # A bare TimeoutError() (== asyncio.TimeoutError since Python 3.11) stringifies
        # to "" -- exactly the shape of error Telethon/httpx can raise on a slow download.
        raise TimeoutError


@pytest.mark.asyncio
async def test_attachment_download_error_with_empty_str_still_yields_valid_error(
    tmp_path: Path,
) -> None:
    """Regression: CollectionError.message requires >=1 character, but str(exc) on a
    bare TimeoutError is "". A failed download must still produce a valid batch --
    with the message text still delivered, just without a download reference --
    instead of raising a secondary pydantic ValidationError that aborts collection
    for every remaining source in that run (see sources/_shared.py:error_message)."""
    message = _FakeMessage(
        id=502,
        date=datetime(2026, 8, 31, 7, 0, tzinfo=UTC),
        message="Ещё один отчёт",
        document=object(),
        file=_FakeFile(
            name="report.docx",
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            size=12345,
            ext=".docx",
        ),
    )
    client = _FailingDownloadTelethonClient([message])
    connector = TelethonConnector(client, download_dir=tmp_path / "downloads")
    batch = await connector.collect(expertosphere_source(), CollectionCursor(message_id=500))

    assert batch.next_cursor == CollectionCursor(message_id=502)
    assert len(batch.errors) == 1
    assert batch.errors[0].code == "attachment_download_failed"
    assert batch.errors[0].message == "TimeoutError"
    assert "Ещё один отчёт" in (batch.raw_items[0].raw_text or "")
    assert batch.raw_items[0].attachments[0].download_ref is None
