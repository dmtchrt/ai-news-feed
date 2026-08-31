"""Local PDF/DOCX extraction; OCR is intentionally outside MVP."""

from __future__ import annotations

from pathlib import Path

import pdfplumber
from docx import Document


class DocumentExtractionError(ValueError):
    pass


def extract_document(path: Path, mime_type: str | None = None) -> str:
    suffix = path.suffix.lower()
    normalized_mime = (mime_type or "").split(";", maxsplit=1)[0].strip().lower()
    if suffix == ".pdf" or normalized_mime == "application/pdf":
        return _extract_pdf(path)
    if suffix == ".docx" or normalized_mime == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ):
        return _extract_docx(path)
    if suffix in {".txt", ".md"} or normalized_mime.startswith("text/"):
        return _nonempty(path.read_text(encoding="utf-8"), path)
    raise DocumentExtractionError(f"unsupported document type: {mime_type or suffix or 'unknown'}")


def _extract_pdf(path: Path) -> str:
    pages: list[str] = []
    try:
        with pdfplumber.open(path) as document:
            for page in document.pages:
                text = page.extract_text(x_tolerance=2, y_tolerance=2)
                if text and text.strip():
                    pages.append(text.strip())
    except Exception as exc:  # pdfplumber wraps several parser exception types.
        raise DocumentExtractionError(f"cannot read PDF {path.name}: {exc}") from exc
    if not pages:
        raise DocumentExtractionError(
            f"PDF {path.name} contains no extractable text (a scanned PDF needs OCR)"
        )
    return "\n\n".join(pages)


def _extract_docx(path: Path) -> str:
    try:
        document = Document(str(path))
    except Exception as exc:  # python-docx delegates ZIP/XML errors.
        raise DocumentExtractionError(f"cannot read DOCX {path.name}: {exc}") from exc
    blocks = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                blocks.append(" | ".join(cells))
    return _nonempty("\n".join(blocks), path)


def _nonempty(text: str, path: Path) -> str:
    cleaned = text.strip()
    if not cleaned:
        raise DocumentExtractionError(f"document {path.name} contains no extractable text")
    return cleaned
